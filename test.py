from docx import Document
import pandas as pd
import os
from os import listdir
from os.path import isfile, join
import numpy as np


def get_text(dir_path):
    document_list = []
    documents = []
    for path, subdirs, files in os.walk(str(dir_path)):
        print('ii')
        for name in files:
            if(os.path.splitext(os.path.join(path, name))[1] == ".docx" or os.path.splitext(os.path.join(path, name))[1] == ".DOCX"):
                document_list.append(os.path.join(path, name))
                document = Document(os.path.join(path, name))
                documents.append(document)


    print(len(documents))
    print(len(files))
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    import re

    final_list = []
    documents_word_list =[]
    resume_text = []
    for i_documents in range(0, len(documents)):

        paragraph = documents[i_documents].paragraphs
        sample_list = []

        for i in paragraph:
            #print(i.text)
            sample_list.append(i.text)

        #print('***', i_documents)
        #print('number of tables : ', len(documents[i_documents].tables))
        if(len(documents[i_documents].tables) > 0):
            for i_table_count in range(0, len(documents[i_documents].tables)):
                table = documents[i_documents].tables[i_table_count]
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            sample_list.append(paragraph.text)

        #print('\n\n\nLIST')
        #print(sample_list)

        for i_sample_list in range(0, len(sample_list) - 1):
            if((sample_list[i_sample_list]).startswith(' Mobile:')):
                del sample_list[i_sample_list]
        #print(sample_list)

        sample_list = ' '.join(sample_list)
        #print(sample_list)

        stop_words = set(stopwords.words('english'))

        word_tokens = word_tokenize(sample_list)

        filtered_sentence = [w for w in word_tokens if not w in stop_words]

        filtered_sentence = []

        for w in word_tokens:
            if w not in stop_words:
                filtered_sentence.append(w)

    #     print(len(word_tokens))
    #     print('\nAfter filtering\n')
    #     print(len(filtered_sentence))

        filtered_sentence = ' '.join(filtered_sentence)
        final = []
        for k in filtered_sentence.split("\n"):
            final = " ".join(re.findall(r"[a-zA-Z0-9]+", k))


    #     print('\nAfter removing special characters\n')
        final = final.split()
    #     print(len(final))

        final = unique_list(final)
#         print(final)
        final_list.append([files[i_documents],final])

    return final_list

def unique_list(l):
    ulist = []
    [ulist.append(x) for x in l if x not in ulist]
    return ulist

def csvMod(dir_path,write_type,fname):

    resume_text = get_text(dir_path)
    print(len(resume_text))
    final_skills = readSkills('all_skills.txt')
    final_skills=sorted(final_skills, key=lambda s: s.lower())
    print(final_skills[:10])
    sorted_index = alphabet_index_list(final_skills)
    print('index',sorted_index)

    matched = []
    # loop for all files:
    for i in range(len(resume_text)):
        # print(resume_text[i][1])
        matched.append(match(final_skills,sorted_index,resume_text[i][1]))      #final = list of words for every resume

    print(len(matched))
    import pandas as pd
    files =[]
    for i in resume_text:
        files.append(i[0])
    # print(len(files),len(matched),len(known_skills))
    if write_type == 1:
        raw_data = {'filename':files,
                    'skills':matched,
                }

        df = pd.DataFrame(raw_data, columns = ['filename', 'skills'])
        df.to_csv(str(fname))
    elif write_type == 2:
        print('Not ready ..')
        # return matched

def alphabet_index_list(l):
    cl = l
    temp = ''
    t = 0
    index_list = np.full((26),-1)
    for i in range(0 , len(cl)):
#         print(i,cl[i])
        if cl[i][0].isdigit() ==False:
            if(i == 0):
                temp = cl[i][0].lower()
                t = ord(temp)
                index_list[t - 97] = i
                continue

            if(cl[i][0].lower() == temp):
                continue
            else:
                temp = cl[i][0].lower()
                t = ord(temp)
                index_list[t - 97] = i
    return index_list


def match(a_sorted,index,r_list):
    matched_words = []
    for i in r_list:
        if i[0].isdigit() == False :
            # print(i[0])
            t = ord(i[0].lower())        # ascii value for first letter
            start = index[t-97]
            if start == -1:
                continue
            else:
                stop = len(index)
                for j in range(t-96,len(index)):
                    if index[j] != -1:
                        stop = index[j]
                        break
            if t < 97:
                start = 0
                stop = index[0]
                break
            for j in a_sorted[start:stop]:
                if i.lower() == j.lower():
                    matched_words.append(i)
                    break
    return matched_words

def readSkills(filename):
    with open(filename) as f:
        l = f.read().splitlines()
    print(len(l),'lines in file',filename)
    return l

