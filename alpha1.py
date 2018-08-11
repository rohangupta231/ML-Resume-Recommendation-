
from docx import Document
import pandas as pd
import os
from os import listdir
from os.path import isfile, join
import numpy as np
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import re

def description_manipulation(directory,read_type):
    file_path = (str)(directory)
    if read_type == 0:
        document = Document(file_path)

        sample_list = []
        paragraph = document.paragraphs
        for i_paragraph in paragraph:
            sample_list.append(i_paragraph.text)

        if(len(document.tables) > 0):
            for i_table_count in range(0, len(document.tables)):
                table = document.tables[i_table_count]
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            sample_list.append(paragraph.text)

            sample_list = ' '.join(sample_list)
    else:
        fp = open(file_path,'r')
        sample_list = fp.read()
        fp.close()

    print('sample_list',sample_list)
    stop_words = set(stopwords.words('english'))
    word_tokens = word_tokenize(sample_list)
    filtered_sentence = [w for w in word_tokens if not w in stop_words]
    filtered_sentence = []

    for w in word_tokens:
        if w not in stop_words:
            filtered_sentence.append(w)

    filtered_sentence = ' '.join(filtered_sentence)
    final = []

    for k in filtered_sentence.split("\n"):
        final = " ".join(re.findall(r"[a-zA-Z0-9]+", k))

    final = final.split()

    final = unique_list(final)
    return final

def sort_list(list1, list2):
    zipped_pairs = zip(list2, list1)
    z = [x for _, x in sorted(zipped_pairs)]
    return z

def match(a_sorted,index,r_list):
    matched_words = []
    for i in r_list:
        if i[0].isdigit() == False :
            t = ord(i[0].lower())        # ascii value for first letter
            start = index[t-97]
            if start == -1:
                continue
            else:
                stop = len(index)
                for j in range(t-96,len(index)):
                    if index[j] != -1:
                        stop = index[j]
                        break
            if t < 97:
                start = 0
                stop = index[0]
                break
            for j in a_sorted[start:stop]:
                if i.lower() == j.lower():
                    matched_words.append(i)
                    break
    print("matched_words",matched_words)
    return matched_words

def matchSkills(l,r):
    matched = []
    for i in l:
        if i in r:
            matched.append(i)
    return matched

def alphabet_index_list(l):
    cl = l
    temp = ''
    t = 0
    index_list = np.full((26),-1)
    for i in range(0 , len(cl)):
        # print(i,cl[i])
        if cl[i][0].isdigit() ==False:
            if(i == 0):
                temp = cl[i][0].lower()
                t = ord(temp)
                index_list[t - 97] = i
                continue

            if(cl[i][0].lower() == temp):
                continue
            else:
                temp = cl[i][0].lower()
                t = ord(temp)
                index_list[t - 97] = i
    print("===========================")
    for i in index_list:
        print(cl[i])
    return index_list

def unique_list(l):
    ulist = []
    [ulist.append(x) for x in l if x not in ulist]
    return ulist

def readSkills(filename):
    with open(filename) as f:
        l = f.read().splitlines()
    print(len(l),'lines in file',filename)
    return l

def write2file(filename, data):
    fp = open(filename,'w')
    fp.write(str(len(data)))
    fp.write('\n')
    for i in range(len(data)):
        fp.write(','.join(data[i]))
        fp.write('\n')
    fp.close()
    print(i,'lines written in file',filename)
